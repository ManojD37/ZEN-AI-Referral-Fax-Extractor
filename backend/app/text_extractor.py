# app/text_extractor.py
import os
from typing import List, Dict, Any
from pathlib import Path
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import docx
from app.log import logger

def extract_text_from_pdf(pdf_path: str, dpi: int = 300) -> str:
    """Extract text from PDF using OCR."""
    logger.info(f"Extracting text from PDF: {pdf_path}")
    try:
        images = convert_from_path(pdf_path, dpi=dpi)
        all_text = []
        
        for i, img in enumerate(images, start=1):
            logger.info(f"Processing page {i}/{len(images)}")
            text = pytesseract.image_to_string(img)
            all_text.append(f"--- Page {i} ---\n{text}")
        
        full_text = "\n\n".join(all_text)
        logger.info(f"PDF extraction complete. Total characters: {len(full_text)}")
        return full_text
    except Exception as e:
        logger.exception(f"PDF extraction failed: {e}")
        raise

def extract_text_from_image(image_path: str) -> str:
    """Extract text from image using OCR."""
    logger.info(f"Extracting text from image: {image_path}")
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img)
        logger.info(f"Image extraction complete. Total characters: {len(text)}")
        return text
    except Exception as e:
        logger.exception(f"Image extraction failed: {e}")
        raise

def extract_text_from_txt(txt_path: str) -> str:
    """Extract text from TXT file."""
    logger.info(f"Reading text file: {txt_path}")
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            text = f.read()
        logger.info(f"Text file read complete. Total characters: {len(text)}")
        return text
    except UnicodeDecodeError:
        # Try with different encoding
        with open(txt_path, 'r', encoding='latin-1') as f:
            text = f.read()
        logger.info(f"Text file read complete (latin-1). Total characters: {len(text)}")
        return text
    except Exception as e:
        logger.exception(f"Text file reading failed: {e}")
        raise

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from Word document (.docx)."""
    logger.info(f"Extracting text from Word document: {docx_path}")
    try:
        doc = docx.Document(docx_path)
        
        # Extract paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Extract tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    tables_text.append(row_text)
        
        # Combine all text
        all_text = "\n".join(paragraphs)
        if tables_text:
            all_text += "\n\n--- Tables ---\n" + "\n".join(tables_text)
        
        logger.info(f"Word document extraction complete. Total characters: {len(all_text)}")
        return all_text
    except Exception as e:
        logger.exception(f"Word document extraction failed: {e}")
        raise

def extract_text_from_file(file_path: str) -> str:
    """
    Universal text extractor - detects file type and extracts text.
    
    Supported formats:
    - PDF: OCR-based extraction
    - Images (JPG, PNG, JPEG): OCR-based extraction
    - TXT: Direct text reading
    - DOCX: Word document text extraction
    """
    file_path = Path(file_path)
    ext = file_path.suffix.lower()
    
    logger.info(f"Starting text extraction for: {file_path} (type: {ext})")
    
    if ext == '.pdf':
        return extract_text_from_pdf(str(file_path))
    elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
        return extract_text_from_image(str(file_path))
    elif ext == '.txt':
        return extract_text_from_txt(str(file_path))
    elif ext == '.docx':
        return extract_text_from_docx(str(file_path))
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def extract_text_with_metadata(file_path: str) -> Dict[str, Any]:
    """
    Extract text along with metadata.
    
    Returns:
        {
            "raw_text": str,
            "file_type": str,
            "character_count": int,
            "word_count": int
        }
    """
    text = extract_text_from_file(file_path)
    
    return {
        "raw_text": text,
        "file_type": Path(file_path).suffix.lower(),
        "character_count": len(text),
        "word_count": len(text.split())
    }